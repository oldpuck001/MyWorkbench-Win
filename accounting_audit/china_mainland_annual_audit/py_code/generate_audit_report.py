# audit_report_generate.py

import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def generate(software_folder, audit_opinion, report_number, report_date):

    basic_info = os.path.join(software_folder, 'data_floder', 'basic_info.csv')
    df = pd.read_csv(basic_info)
    name = df['企业名称'][0]
    period = df['审计期间'][0]
    deadline = df['审计截止日'][0]

    report_floder = os.path.join(software_folder, 'report')
    # 确保保存路径存在
    if not os.path.exists(report_floder):
        os.makedirs(report_floder, exist_ok=True)

    if audit_opinion == 'standard':
        base_dir = os.path.dirname(__file__)
        file_path = os.path.join(base_dir, 'audit_report_standard.txt')
    else:
        pass

    # 格式化日期为 'yyyy年m月d日' 的格式
    formatted_report_date = report_date[:4] + '年' + report_date[5:7] + '月' + report_date[8:] + '日'

    # 定义要替换的旧内容和新内容
    replacement_data = {
        '***公司名称***': name,
        '***报告文号***': report_number,
        '***截止日***': deadline,
        '***期间***': period,
        '***出具日期***': formatted_report_date
    }

    with open(file_path, 'r', encoding='utf-8') as file:
        lines = [line.rstrip('\n') for line in file.readlines()]     # 读取所有行并去除换行符

    # 用于存储修改后的内容
    modified_lines = []
    # 遍历每一行进行替换
    for line in lines:
        # 对每行中的特定内容进行替换
        for old_content, new_content in replacement_data.items():
            line = line.replace(old_content, new_content)
        # 将修改后的行添加到新列表中
        modified_lines.append(line)

    # 创建文档对象
    doc = Document()

    # 报告文号的正则表达式
    report_number_pattern = r'\S*字\S*号$'

    # 将修改后的内容打印出来（或保存到新文件中）
    for modified_line in modified_lines:
        paragraph = doc.add_paragraph()                             # 添加第一个段落并设置字体和对齐方式
        run = paragraph.add_run(modified_line)        
        run.font.name = 'SimSun'                                    # 设置字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')     # 设置中文字体
        if modified_line == '审 计 报 告' or modified_line == '':
            run.font.size = Pt(25)                                  # 设置字号
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     # 设置居中对齐
        elif re.search(report_number_pattern, modified_line):
            run.font.size = Pt(12)                                  # 设置字号
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     # 设置居中对齐
        else:
            run.font.size = Pt(12)                                  # 设置字号
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT       # 设置靠左对齐

    # 添加页眉
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "XXX会计师事务所"
    header_run = header_paragraph.runs[0]
    header_run.font.name = 'SimSun'
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    header_run.font.size = Pt(12)
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE  # 设置分散对齐

    # 添加页脚
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "地址：……          联系电话：***-********"
    footer_run = footer_paragraph.runs[0]
    footer_run.font.name = 'SimSun'
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    footer_run.font.size = Pt(12)
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 保存文档
    save_path = os.path.join(report_floder, '审计报告正文.docx')

    doc.save(save_path)

    return save_path